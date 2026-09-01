from __future__ import annotations

import io
from typing import Any, Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from app.datasets.storage import DatasetRecord
from app.reports.report_content import (
    collect_report_content,
    combined_limitations,
    report_filename,
    safe_text,
)

_NAVY = "172554"
_BLUE = "2563EB"
_LIGHT_BLUE = "E8F0FE"
_LIGHT_GRAY = "F1F5F9"
_WHITE = "FFFFFF"


def build_word_report(record: DatasetRecord) -> bytes:
    """Build a deterministic, styled Word report without provider access."""
    content = collect_report_content(record)
    report = content.report
    document = Document()
    _configure_document(document)
    _add_title(document, report)
    _add_overview(document, report)
    _add_findings(document, report, content.analysis)
    _add_anomalies(document, report.get("anomalies", []))
    _add_correlations(document, report.get("correlations", []))
    _add_caveats(document, content)
    _add_fact_ledger(document, content.analysis)
    _add_limitations(document, combined_limitations(content))
    _add_footer(document)
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def word_report_filename(record: DatasetRecord) -> str:
    return report_filename(record, "docx")


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for name, size in (("Title", 26), ("Heading 1", 17), ("Heading 2", 13)):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(_NAVY if name != "Heading 2" else _BLUE)
        style.paragraph_format.space_before = Pt(16 if name == "Heading 1" else 12)
        style.paragraph_format.space_after = Pt(8 if name == "Heading 1" else 6)


def _add_title(document: Document, report: dict[str, Any]) -> None:
    title = document.add_heading("AI Data Analyst", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle = document.add_paragraph("Deterministic analysis report")
    subtitle.style = document.styles["Subtitle"]
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(_BLUE)
    metadata = document.add_paragraph()
    metadata.add_run(safe_text(report["filename"])).bold = True
    metadata.add_run(f"  |  Generated {safe_text(report['generated_at'])}")
    _paragraph_rule(metadata, _BLUE)


def _add_overview(document: Document, report: dict[str, Any]) -> None:
    document.add_heading("Executive overview", level=1)
    overview = report["overview"]
    rows = [
        ("Rows", f"{overview['rows']:,}"),
        ("Columns", f"{overview['columns']:,}"),
        ("Missing values", f"{overview['missing_total']:,}"),
        ("Duplicate rows", f"{overview['duplicate_rows']:,}"),
    ]
    _table(document, ["Measure", "Value"], rows, widths=(4.7, 1.8))


def _add_findings(document: Document, report: dict[str, Any], analysis: Any | None) -> None:
    document.add_heading("Key findings", level=1)
    for finding in report.get("key_findings", []):
        document.add_paragraph(safe_text(finding), style="List Bullet")
    document.add_heading("Latest grounded analysis", level=2)
    if analysis is None:
        _callout(document, "No Deep Reasoning analysis is available in the current server session.")
    else:
        document.add_paragraph(safe_text(analysis.final_answer_text))


def _add_anomalies(document: Document, anomalies: list[dict[str, Any]]) -> None:
    document.add_heading("Anomalies", level=1)
    if not anomalies:
        document.add_paragraph("No anomaly results are available.")
        return
    rows = [(item.get("column", ""), item.get("anomaly_count", ""), f"{item.get('anomaly_pct', '')}%") for item in anomalies[:25]]
    _table(document, ["Column", "Outliers", "Share"], rows, widths=(3.7, 1.3, 1.5))


def _add_correlations(document: Document, correlations: list[dict[str, Any]]) -> None:
    document.add_heading("Notable correlations", level=1)
    if not correlations:
        document.add_paragraph("No correlation results are available.")
        return
    rows = [
        (item.get("column_a", ""), item.get("column_b", ""), _format_number(item.get("correlation")))
        for item in correlations[:20]
    ]
    _table(document, ["Variable A", "Variable B", "Correlation"], rows, widths=(2.65, 2.65, 1.2))
    _callout(document, "Correlation is descriptive evidence and does not establish causation.")


def _add_caveats(document: Document, content: Any) -> None:
    document.add_heading("Data caveats", level=1)
    rows: list[tuple[Any, ...]] = [
        ("Duplicates", "Dataset", content.caveats.duplicate_row_count, f"{content.caveats.duplicate_pct}%"),
        ("Rows dropped", "Dataset", content.caveats.rows_dropped, content.caveats.rows_dropped_note),
    ]
    rows.extend(("Coverage", item.column, item.non_null_rows, f"{item.coverage_pct}%") for item in content.caveats.column_coverage)
    rows.extend(("Type anomaly", "Dataset", "", item) for item in content.caveats.type_anomalies)
    rows.extend(("Ingestion", notice.column or "Dataset", notice.code, notice.message) for notice in content.ingestion_notices)
    _table(document, ["Category", "Field", "Value", "Details"], rows[:100], widths=(1.3, 1.4, 1.0, 2.8))


def _add_fact_ledger(document: Document, analysis: Any | None) -> None:
    document.add_heading("Fact ledger", level=1)
    facts = list(analysis.fact_ledger.facts) if analysis is not None and analysis.fact_ledger else []
    if not facts:
        document.add_paragraph("No grounded facts are available for this export.")
        return
    rows = [(fact.id, fact.tool, fact.result_path, fact.value, fact.unit or "") for fact in facts[:100]]
    _table(document, ["Fact", "Tool", "Result", "Value", "Unit"], rows, widths=(1.1, 1.4, 2.0, 1.2, 0.8))


def _add_limitations(document: Document, limitations: list[Any]) -> None:
    document.add_heading("Limitations", level=1)
    if not limitations:
        document.add_paragraph("No limitations were identified for this report.")
        return
    for item in limitations:
        paragraph = document.add_paragraph(style="List Bullet")
        run = paragraph.add_run(f"{safe_text(item.severity)} — {safe_text(item.category)}: ")
        run.bold = True
        paragraph.add_run(safe_text(item.text))


def _table(document: Document, headers: list[str], rows: Iterable[tuple[Any, ...]], widths: tuple[float, ...]) -> None:
    rows = list(rows)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    table.style = "Table Grid"
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)
    _set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.width = Inches(widths[index])
        cell.text = safe_text(header)
        _shade(cell, _BLUE)
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string(_WHITE)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row_number, values in enumerate(rows, start=1):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].width = Inches(widths[index])
            cells[index].text = safe_text(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_number % 2 == 0:
                _shade(cells[index], _LIGHT_GRAY)
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.space_before = Pt(2)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _callout(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(safe_text(text))
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    properties = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), _LIGHT_BLUE)
    properties.append(shading)


def _shade(cell: Any, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), color)


def _set_table_geometry(table: Any, widths: tuple[float, ...]) -> None:
    properties = table._tbl.tblPr
    table_width = properties.first_child_found_in("w:tblW")
    table_width.set(qn("w:type"), "dxa")
    table_width.set(qn("w:w"), "9360")
    indent = properties.first_child_found_in("w:tblInd")
    if indent is None:
        indent = OxmlElement("w:tblInd")
        properties.append(indent)
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    grid = table._tbl.tblGrid
    for grid_column, width in zip(grid.gridCol_lst, widths):
        grid_column.set(qn("w:w"), str(round(width * 1440)))


def _paragraph_rule(paragraph: Any, color: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    properties.append(borders)


def _add_footer(document: Document) -> None:
    for section in document.sections:
        paragraph = section.footer.paragraphs[0]
        paragraph.text = "AI Data Analyst | Deterministic report export"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.runs[0].font.size = Pt(8)
        paragraph.runs[0].font.color.rgb = RGBColor(100, 116, 139)


def _format_number(value: Any) -> str:
    try:
        return f"{float(value):.3f}"
    except (TypeError, ValueError):
        return safe_text(value)
